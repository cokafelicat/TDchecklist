#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
from typing import List, Dict, Set, Optional
import PyPDF2
from docx import Document
import logging
from pathlib import Path
from src.database import DatabaseManager
import argparse

class DocumentAnalyzer:
    def __init__(self, keywords_file: Optional[str] = None):
        """初始化文档分析器
        keywords_file: 关键词配置文件路径（可选，暂未实现文件加载，仅为兼容参数）
        """
        self.logger = logging.getLogger(__name__)
        self.db = DatabaseManager()
        self._refresh_keywords()
    
    def _refresh_keywords(self):
        """刷新关键词列表"""
        self.keywords = {kw["keyword"] for kw in self.db.get_all_keywords()}
    
    def clean_keyword(self, keyword: str) -> str:
        """清理关键词，去除引号和多余的空格"""
        return keyword.strip().strip('"').strip("'").strip()

    def add_keywords(self, keywords: List[str]) -> None:
        """添加关键词到数据库中"""
        for keyword in keywords:
            cleaned_keyword = self.clean_keyword(keyword)
            if cleaned_keyword:
                self.db.add_keyword(cleaned_keyword)
        self._refresh_keywords()

    def remove_keywords(self, keywords: List[str]) -> None:
        """从数据库中删除关键词"""
        for keyword in keywords:
            cleaned_keyword = self.clean_keyword(keyword)
            if cleaned_keyword:
                self.db.delete_keyword(cleaned_keyword)
        self._refresh_keywords()

    def get_keywords(self) -> Set[str]:
        """获取所有关键词"""
        return self.keywords

    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, any]]:
        """从PDF文件中提取文本
        Returns:
            包含页码和文本内容的字典列表
        """
        pages_content = []
        # 打开PDF文件
        with open(pdf_path, 'rb') as file:
            # 创建PDF阅读器
            pdf_reader = PyPDF2.PdfReader(file)
            # 遍历每一页
            for page_num, page in enumerate(pdf_reader.pages, 1):
                # 提取文本并添加到列表中
                text = page.extract_text()
                if text.strip():
                    pages_content.append({
                        'page': page_num,
                        'content': text
                    })
        return pages_content

    def extract_text_from_docx(self, docx_path: str) -> List[Dict[str, any]]:
        """从DOCX文件中提取文本
        Returns:
            包含段落位置和文本内容的字典列表
        """
        # 打开DOCX文件
        doc = Document(docx_path)
        paragraphs_content = []
        current_position = 1
        
        # 提取正文段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs_content.append({
                    'page': current_position,
                    'content': paragraph.text.strip()
                })
                current_position += 1
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs_content.append({
                        'page': current_position,
                        'content': row_text
                    })
                    current_position += 1
                    
        return paragraphs_content

    def truncate_text(self, text: str, max_length: int = 100) -> str:
        """截断文本，保持完整的中文句子
        Args:
            text: 要截断的文本
            max_length: 最大长度
        Returns:
            截断后的文本
        """
        if len(text) <= max_length:
            return text
            
        # 在最大长度位置寻找句号、问号或感叹号
        cutoff = max_length
        punctuations = ['。', '？', '！', '.', '?', '!']
        
        # 从最大长度位置向前查找标点符号
        while cutoff > max_length // 2:
            if text[cutoff - 1] in punctuations:
                return text[:cutoff] + '...'
            cutoff -= 1
            
        # 如果找不到合适的断句点，直接截断
        return text[:max_length] + '...'

    def extract_section_number(self, text: str) -> str:
        """提取段落的章节编号
        Args:
            text: 段落文本
        Returns:
            章节编号，如果没有则返回空字符串
        """
        # 匹配常见的章节编号格式
        import re
        patterns = [
            r'^第[一二三四五六七八九十百零]+章',  # 中文数字章节
            r'^第\d+章',                      # 阿拉伯数字章节
            r'^\d+\.\d+(\.\d+)?',            # 数字编号（如1.2, 1.2.3）
            r'^[一二三四五六七八九十]+、',      # 中文数字编号
            r'^\d+、'                         # 阿拉伯数字编号
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return match.group(0)
        return ""

    def find_relevant_paragraphs(self, pages_content: List[Dict[str, any]]) -> List[Dict[str, any]]:
        """查找包含关键词的相关句子（前后两个句号之间），并高亮该句子。优化为真正提取前后两个句号之间的内容。"""
        import re
        results = []
        current_section = ""
        # 支持中英文句号
        sentence_end = r'[。！？.!?]'
        for page_data in pages_content:
            page_num = page_data['page']
            content = page_data['content']
            # 先查找所有关键词在全文的位置
            lowered_content = content.lower()
            for keyword in self.keywords:
                keyword_lower = keyword.lower()
                start = 0
                while True:
                    idx = lowered_content.find(keyword_lower, start)
                    if idx == -1:
                        break
                    # 向前找最近的句号
                    pre = content.rfind('。', 0, idx)
                    for p in ['。', '！', '？', '.', '!', '?']:
                        p_idx = content.rfind(p, 0, idx)
                        if p_idx > pre:
                            pre = p_idx
                    # 向后找下一个句号
                    post = len(content)
                    for p in ['。', '！', '？', '.', '!', '?']:
                        p_idx = content.find(p, idx + len(keyword))
                        if p_idx != -1 and p_idx < post:
                            post = p_idx
                    # 取出完整句子
                    sent = content[pre+1:post+1].strip()
                    # 提取章节号
                    section = self.extract_section_number(sent)
                    if section:
                        current_section = section
                    # 避免重复添加同一句
                    already = False
                    for r in results:
                        if r['page'] == page_num and r['text'] == sent:
                            already = True
                            break
                    if not already and sent:
                        results.append({
                            'page': page_num,
                            'section': current_section,
                            'text': sent,
                            'keyword': keyword,
                            'original_length': len(sent)
                        })
                    start = idx + len(keyword)
        return results

    def process_document(self, file_path: str) -> List[Dict[str, any]]:
        """处理文档并返回结果"""
        # 判断文件格式
        if file_path.lower().endswith('.pdf'):
            # 如果是PDF文件，提取文本
            pages_content = self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            # 如果是DOCX文件，提取文本
            pages_content = self.extract_text_from_docx(file_path)
        else:
            # 如果不是支持的文件格式，抛出异常
            raise ValueError("不支持的文件格式，仅支持PDF和DOCX格式")
        
        # 返回结果
        return self.find_relevant_paragraphs(pages_content)

def highlight_keywords(text: str, keywords: Set[str], color: str = '\033[93m', bold: bool = True) -> str:
    """高亮并加粗显示文本中的关键词所在句子
    Args:
        text: 要处理的句子
        keywords: 关键词集合
        color: ANSI转义序列颜色代码，默认为黄色
        bold: 是否加粗
    Returns:
        处理后的文本，关键词会被高亮加粗显示
    """
    reset_color = '\033[0m'
    bold_code = '\033[1m' if bold else ''
    result = text
    for keyword in sorted(keywords, key=len, reverse=True):
        pattern = keyword.lower()
        index = result.lower().find(pattern)
        while index != -1:
            original_keyword = result[index:index + len(keyword)]
            result = (
                result[:index] +
                f"{bold_code}{color}{original_keyword}{reset_color}" +
                result[index + len(keyword):]
            )
            next_start = index + len(keyword) + len(color) + len(reset_color) + (len(bold_code) if bold else 0)
            index = result.lower().find(pattern, next_start)
    # 整句加粗
    if bold:
        result = f"{bold_code}{result}{reset_color}"
    return result

def print_checklist(results: List[Dict[str, any]], keywords: Set[str]):
    """以checklist格式打印结果，并高亮显示关键词
    Args:
        results: 匹配结果列表
        keywords: 关键词集合
    """
    if not results:
        print("\n未找到任何匹配的段落。")
        return

    yellow = '\033[93m'   # 黄色
    cyan = '\033[96m'     # 青色
    green = '\033[92m'    # 绿色
    gray = '\033[90m'     # 灰色
    reset = '\033[0m'     # 重置颜色
    bold = '\033[1m'      # 粗体
    
    print(f"\n{bold}找到 {len(results)} 个匹配段落：{reset}")
    print(f"{cyan}{'-' * 65}{reset}")
    
    for i, result in enumerate(results, 1):
        # 高亮显示段落中的所有关键词
        highlighted_text = highlight_keywords(result['text'], keywords)
        
        # 构建位置信息
        location = f"第{result['page']}页"
        if result['section']:
            location = f"{result['section']} | {location}"
            
        # 显示段落信息
        print(f"{cyan}[{i}]{reset} {green}{location}{reset}")
        
        # 显示正文内容
        print(f"{highlighted_text}")
        
        # 如果原文比较长，显示省略信息
        if result.get('original_length', 0) > 200:
            remaining = result['original_length'] - 200
            print(f"{gray}... 已省略后续 {remaining} 字{reset}")
            
        print(f"{cyan}{'-' * 65}{reset}\n")

def main():
    parser = argparse.ArgumentParser(description='文档关键词段落提取工具')
    parser.add_argument('file', nargs='?', help='要处理的文档路径（PDF或DOCX格式）')
    parser.add_argument('--add-keywords', nargs='+', help='添加新的关键词')
    parser.add_argument('--remove-keywords', nargs='+', help='删除指定的关键词')
    parser.add_argument('--remove-all', action='store_true', help='删除所有关键词')
    parser.add_argument('--list-keywords', action='store_true', help='列出所有当前的关键词')
    parser.add_argument('--keywords-file', default='keywords.json', help='指定关键词配置文件路径')
    
    args = parser.parse_args()
    
    analyzer = DocumentAnalyzer(keywords_file=args.keywords_file)
    
    # 处理关键词相关的命令
    if args.list_keywords:
        print("\n当前的关键词列表：")
        for keyword in sorted(analyzer.keywords):
            print(f"- {keyword}")
        return
        
    if args.add_keywords:
        cleaned_keywords = [analyzer.clean_keyword(k) for k in args.add_keywords]
        analyzer.add_keywords(args.add_keywords)
        print(f"\n已添加关键词: {', '.join(cleaned_keywords)}")
        
    if args.remove_all:
        # 保存关键词数量用于显示
        removed_count = len(analyzer.keywords)
        # 清空所有关键词
        analyzer.keywords.clear()
        analyzer.save_keywords([])
        print(f"\n已删除所有关键词（共 {removed_count} 个）")
    elif args.remove_keywords:
        cleaned_keywords = [analyzer.clean_keyword(k) for k in args.remove_keywords]
        analyzer.remove_keywords(args.remove_keywords)
        print(f"\n已删除关键词: {', '.join(cleaned_keywords)}")
    
    # 如果没有提供文件参数且没有其他操作，显示帮助信息
    if not any([args.file, args.add_keywords, args.remove_keywords, args.list_keywords]):
        parser.print_help()
        return
    
    # 处理文档
    if args.file:
        if not os.path.exists(args.file):
            print(f"\n错误：找不到文件 {args.file}")
            return
            
        if not analyzer.keywords:
            print("\n错误：未设置任何关键词。请先使用 --add-keywords 添加关键词。")
            return
            
        try:
            results = analyzer.process_document(args.file)
            print_checklist(results, analyzer.keywords)
        except Exception as e:
            print(f"\n处理文档时出错: {str(e)}")

if __name__ == "__main__":
    main()
